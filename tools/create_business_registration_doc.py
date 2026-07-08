from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "business_registration_document"
DOCX_PATH = OUT_DIR / "Nico_Online_Mart_Business_Registration_Document.docx"
PDF_PATH = OUT_DIR / "Nico_Online_Mart_Business_Registration_Document.pdf"
LOGO_PATH = ROOT / "assets" / "images" / "logo.png"


NAVY = "092044"
GOLD = "B98A3A"
GREEN = "1F8A4C"
LIGHT = "F4F7FA"
MID = "D9E2EC"
TEXT = "1F2937"


def hex_color(value):
    return colors.HexColor(f"#{value}")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_padding(cell, top=120, left=160, bottom=120, right=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = tc_pr.first_child_found_in("w:tcMar")
    if margin is None:
        margin = OxmlElement("w:tcMar")
        tc_pr.append(margin)
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = margin.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margin.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=None, bold=False, color=TEXT, italic=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    style_run(run, size=15 if level == 1 else 12, bold=True, color=NAVY)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    style_run(run, size=9.6, color=TEXT)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.18)
        p.paragraph_format.space_after = Pt(3.5)
        r = p.add_run("- ")
        style_run(r, size=9.4, color=GREEN, bold=True)
        r = p.add_run(item)
        style_run(r, size=9.4, color=TEXT)


def add_key_table(doc):
    rows = [
        ("Application / Trading Name", "Nico Online Mart / Nico Mart Sri Lanka"),
        ("Business Activity", "Online retail, grocery, bottled water, food ordering, and home/office delivery support."),
        ("Software Type", "Mobile commerce application built with Flutter and Firebase services."),
        ("Customer Channel", "Mobile application and related web/contact channels."),
        ("Public Contact Listed in App", "WhatsApp: +94 329 013 2841; Website: www.nicomart.it"),
        ("Legal Owner / Registered Entity", "[To be completed by applicant]"),
        ("Registered Business Address", "[To be completed by applicant]"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5.3)
    table.columns[1].width = Cm(10.2)
    hdr = table.rows[0].cells
    hdr[0].text = "Registration Detail"
    hdr[1].text = "Information"
    for cell in hdr:
        shade_cell(cell, NAVY)
        set_cell_border(cell, NAVY)
        set_cell_padding(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                style_run(r, size=9.2, bold=True, color="FFFFFF")
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        shade_cell(cells[0], LIGHT)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_padding(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    style_run(r, size=9.2, bold=(cell is cells[0]), color=TEXT)
    doc.add_paragraph()


def add_signature_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("Applicant / Authorized Person", "[Name]"),
        ("Designation", "[Owner / Director / Authorized Representative]"),
        ("Signature", "______________________________"),
        ("Date", "01 May 2026"),
    ]
    for row, (left, right) in zip(table.rows, data):
        row.cells[0].text = left
        row.cells[1].text = right
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_padding(cell, top=140, bottom=140)
            for p in cell.paragraphs:
                for r in p.runs:
                    style_run(r, size=9.3, bold=(cell is row.cells[0]), color=TEXT)


def build():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.55)
    sec.bottom_margin = Cm(1.55)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(9.6)

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Business Registration Support Document")
    style_run(hr, size=8.5, color="667085")

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Nico Online Mart - Confidential business registration document")
    style_run(fr, size=8, color="667085")

    title_table = doc.add_table(rows=1, cols=2)
    title_table.autofit = False
    title_table.columns[0].width = Cm(4.4)
    title_table.columns[1].width = Cm(11.2)
    left, right = title_table.rows[0].cells
    set_cell_border(left, "FFFFFF")
    set_cell_border(right, "FFFFFF")
    set_cell_padding(left, top=0, left=0, right=100, bottom=60)
    set_cell_padding(right, top=120, left=120, bottom=60)
    if LOGO_PATH.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Inches(1.35))
    rp = right.paragraphs[0]
    rp.paragraph_format.space_after = Pt(2)
    r = rp.add_run("NICO ONLINE MART")
    style_run(r, size=24, bold=True, color=NAVY)
    rp2 = right.add_paragraph()
    r = rp2.add_run("Business Registration App Description and Operational Declaration")
    style_run(r, size=12.2, bold=True, color=GOLD)
    rp3 = right.add_paragraph()
    r = rp3.add_run("Prepared for business registration, licensing, and administrative review")
    style_run(r, size=9.5, color="475467")

    add_key_table(doc)

    add_heading(doc, "1. Purpose of This Document")
    add_body(
        doc,
        "This document describes the Nico Online Mart mobile application and the business activities it supports. "
        "It may be submitted as a supporting description for business registration or licensing review. It is not a replacement for official government forms, tax registrations, licenses, or professional legal advice.",
    )

    add_heading(doc, "2. Nature of the Application")
    add_body(
        doc,
        "Nico Online Mart is a mobile commerce and delivery application for customers to browse products, place orders, and request delivery of bottled water, Sri Lankan grocery items, food items, and related retail products. "
        "The app connects customers with listed products or shops, records order information, and supports customer communication for order assistance.",
    )

    add_heading(doc, "3. Main App Functions")
    add_bullets(
        doc,
        [
            "Customer registration, login, account access, and password recovery.",
            "Product, grocery, bottled water, food, shop, and restaurant browsing with images, pricing, categories, and availability information.",
            "Cart management, quantity updates, discount display, delivery fee calculation, and order summary.",
            "Checkout with customer name, phone number, delivery address, optional delivery instructions, and order confirmation.",
            "Order history, payment information, delivery status, and order tracking screens.",
            "Cash on delivery and online payment support through configured payment providers where activated.",
            "Customer support contact through WhatsApp and related app communication channels.",
            "Location and address support for delivery coordination.",
        ],
    )

    add_heading(doc, "4. Business Activities Supported")
    add_bullets(
        doc,
        [
            "Online sale and delivery coordination for groceries, bottled water, food items, and other approved retail products.",
            "Collection of order details needed to process purchases and deliver goods to customers.",
            "Acceptance of customer payments through approved payment methods, including cash on delivery and enabled online payment gateways.",
            "Operation of a customer support channel for product, payment, delivery, and technical assistance.",
            "Maintenance of digital records for orders, payments, delivery status, and customer service follow-up.",
        ],
    )

    add_heading(doc, "5. Technology and Data Handling")
    add_body(
        doc,
        "The application is built with Flutter and uses Firebase services for authentication, database storage, cloud functions, and media storage. Customer data handled by the app is limited to information required for account access, order processing, delivery coordination, payment confirmation, and support communication.",
    )
    add_bullets(
        doc,
        [
            "Typical data includes customer name, email, phone number, delivery address, order details, payment status, and delivery notes.",
            "Card processing, where enabled, is handled through third-party payment service providers; the app should not store full payment card details.",
            "Access to business data should be limited to authorized personnel for order fulfilment, accounting, support, and compliance purposes.",
        ],
    )

    add_heading(doc, "6. Operational and Legal Responsibilities")
    add_bullets(
        doc,
        [
            "Register the business under the correct legal owner or company name before commercial operation.",
            "Maintain accurate product descriptions, prices, delivery fees, stock status, and customer-facing terms.",
            "Comply with applicable consumer protection, e-commerce, food safety, delivery, tax, and data protection requirements.",
            "Issue receipts or order confirmations and maintain records required for accounting and legal compliance.",
            "Publish clear policies for delivery, cancellation, refunds, privacy, and customer support before public launch.",
            "Use only lawful product images, trademarks, business names, and supplier content.",
        ],
    )

    add_heading(doc, "7. Declaration")
    add_body(
        doc,
        "The applicant declares that Nico Online Mart is intended to operate as a lawful online retail and delivery-support business. The applicant accepts responsibility for completing all official registration forms, obtaining any required permits, maintaining accurate business records, and complying with the laws applicable to the location of operation.",
    )
    add_signature_table(doc)

    doc.save(DOCX_PATH)
    build_pdf()
    return DOCX_PATH, PDF_PATH


def make_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleNico",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=29,
            textColor=hex_color(NAVY),
            alignment=TA_LEFT,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SubtitleNico",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=hex_color(GOLD),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallMuted",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=hex_color("667085"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="HeadingNico",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=hex_color(NAVY),
            spaceBefore=11,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyNico",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.4,
            leading=12.4,
            textColor=hex_color(TEXT),
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletNico",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.1,
            leading=12,
            leftIndent=12,
            bulletIndent=0,
            textColor=hex_color(TEXT),
            spaceAfter=3.2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TableText",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11,
            textColor=hex_color(TEXT),
        )
    )
    styles.add(
        ParagraphStyle(
            name="TableBold",
            parent=styles["TableText"],
            fontName="Helvetica-Bold",
        )
    )
    return styles


def p(text, style):
    return Paragraph(text, style)


def pdf_heading(story, text, styles):
    story.append(Paragraph(text, styles["HeadingNico"]))


def pdf_body(story, text, styles):
    story.append(Paragraph(text, styles["BodyNico"]))


def pdf_bullets(story, items, styles):
    for item in items:
        story.append(Paragraph(item, styles["BulletNico"], bulletText="-"))


def build_pdf():
    styles = make_pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.55 * cm,
        leftMargin=1.55 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.35 * cm,
        title="Nico Online Mart Business Registration Document",
        author="Nico Online Mart",
    )

    story = []
    logo = Image(str(LOGO_PATH), width=3.1 * cm, height=2.35 * cm) if LOGO_PATH.exists() else ""
    title_block = [
        Paragraph("NICO ONLINE MART", styles["TitleNico"]),
        Paragraph("Business Registration App Description and Operational Declaration", styles["SubtitleNico"]),
        Paragraph("Prepared for business registration, licensing, and administrative review", styles["SmallMuted"]),
    ]
    hero = Table([[logo, title_block]], colWidths=[3.6 * cm, 12.4 * cm])
    hero.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
            ]
        )
    )
    story.append(hero)

    details = [
        ["Registration Detail", "Information"],
        ["Application / Trading Name", "Nico Online Mart / Nico Mart Sri Lanka"],
        ["Business Activity", "Online retail, grocery, bottled water, food ordering, and home/office delivery support."],
        ["Software Type", "Mobile commerce application built with Flutter and Firebase services."],
        ["Customer Channel", "Mobile application and related web/contact channels."],
        ["Public Contact Listed in App", "WhatsApp: +94 329 013 2841; Website: www.nicomart.it"],
        ["Legal Owner / Registered Entity", "[To be completed by applicant]"],
        ["Registered Business Address", "[To be completed by applicant]"],
    ]
    detail_table = Table(
        [[p(str(c), styles["TableBold"] if i == 0 or j == 0 else styles["TableText"]) for j, c in enumerate(row)] for i, row in enumerate(details)],
        colWidths=[5.1 * cm, 10.9 * cm],
        hAlign="CENTER",
    )
    detail_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), hex_color(NAVY)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("BACKGROUND", (0, 1), (0, -1), hex_color(LIGHT)),
                ("GRID", (0, 0), (-1, -1), 0.45, hex_color(MID)),
                ("BOX", (0, 0), (-1, -1), 0.8, hex_color(NAVY)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(detail_table)
    story.append(Spacer(1, 0.22 * cm))

    sections = [
        (
            "1. Purpose of This Document",
            [
                "This document describes the Nico Online Mart mobile application and the business activities it supports. It may be submitted as a supporting description for business registration or licensing review. It is not a replacement for official government forms, tax registrations, licenses, or professional legal advice."
            ],
            [],
        ),
        (
            "2. Nature of the Application",
            [
                "Nico Online Mart is a mobile commerce and delivery application for customers to browse products, place orders, and request delivery of bottled water, Sri Lankan grocery items, food items, and related retail products. The app connects customers with listed products or shops, records order information, and supports customer communication for order assistance."
            ],
            [],
        ),
        (
            "3. Main App Functions",
            [],
            [
                "Customer registration, login, account access, and password recovery.",
                "Product, grocery, bottled water, food, shop, and restaurant browsing with images, pricing, categories, and availability information.",
                "Cart management, quantity updates, discount display, delivery fee calculation, and order summary.",
                "Checkout with customer name, phone number, delivery address, optional delivery instructions, and order confirmation.",
                "Order history, payment information, delivery status, and order tracking screens.",
                "Cash on delivery and online payment support through configured payment providers where activated.",
                "Customer support contact through WhatsApp and related app communication channels.",
                "Location and address support for delivery coordination.",
            ],
        ),
        (
            "4. Business Activities Supported",
            [],
            [
                "Online sale and delivery coordination for groceries, bottled water, food items, and other approved retail products.",
                "Collection of order details needed to process purchases and deliver goods to customers.",
                "Acceptance of customer payments through approved payment methods, including cash on delivery and enabled online payment gateways.",
                "Operation of a customer support channel for product, payment, delivery, and technical assistance.",
                "Maintenance of digital records for orders, payments, delivery status, and customer service follow-up.",
            ],
        ),
        (
            "5. Technology and Data Handling",
            [
                "The application is built with Flutter and uses Firebase services for authentication, database storage, cloud functions, and media storage. Customer data handled by the app is limited to information required for account access, order processing, delivery coordination, payment confirmation, and support communication."
            ],
            [
                "Typical data includes customer name, email, phone number, delivery address, order details, payment status, and delivery notes.",
                "Card processing, where enabled, is handled through third-party payment service providers; the app should not store full payment card details.",
                "Access to business data should be limited to authorized personnel for order fulfilment, accounting, support, and compliance purposes.",
            ],
        ),
        (
            "6. Operational and Legal Responsibilities",
            [],
            [
                "Register the business under the correct legal owner or company name before commercial operation.",
                "Maintain accurate product descriptions, prices, delivery fees, stock status, and customer-facing terms.",
                "Comply with applicable consumer protection, e-commerce, food safety, delivery, tax, and data protection requirements.",
                "Issue receipts or order confirmations and maintain records required for accounting and legal compliance.",
                "Publish clear policies for delivery, cancellation, refunds, privacy, and customer support before public launch.",
                "Use only lawful product images, trademarks, business names, and supplier content.",
            ],
        ),
    ]
    for heading, paras, bullets in sections:
        pdf_heading(story, heading, styles)
        for text in paras:
            pdf_body(story, text, styles)
        if bullets:
            pdf_bullets(story, bullets, styles)

    pdf_heading(story, "7. Declaration", styles)
    pdf_body(
        story,
        "The applicant declares that Nico Online Mart is intended to operate as a lawful online retail and delivery-support business. The applicant accepts responsibility for completing all official registration forms, obtaining any required permits, maintaining accurate business records, and complying with the laws applicable to the location of operation.",
        styles,
    )
    sig = [
        ["Applicant / Authorized Person", "[Name]"],
        ["Designation", "[Owner / Director / Authorized Representative]"],
        ["Signature", "______________________________"],
        ["Date", "01 May 2026"],
    ]
    sig_table = Table(
        [[p(c, styles["TableBold"] if j == 0 else styles["TableText"]) for j, c in enumerate(row)] for row in sig],
        colWidths=[5.7 * cm, 10.3 * cm],
    )
    sig_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, hex_color(MID)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(KeepTogether([sig_table]))

    def page_canvas(canvas, doc_obj):
        canvas.saveState()
        canvas.setStrokeColor(hex_color(GOLD))
        canvas.setLineWidth(1.2)
        canvas.line(doc_obj.leftMargin, A4[1] - 1.05 * cm, A4[0] - doc_obj.rightMargin, A4[1] - 1.05 * cm)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(hex_color("667085"))
        canvas.drawCentredString(A4[0] / 2, 0.75 * cm, f"Nico Online Mart - Confidential business registration document - Page {doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=page_canvas, onLaterPages=page_canvas)


if __name__ == "__main__":
    print(build())
